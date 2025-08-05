import os
import json
from typing import List

import matplotlib.pyplot as plt
from datetime import datetime
import numpy as np
from PIL import Image
import tempfile
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import io
import re

from video.ui_utils import load_custom_css
# 移除重复导入，保持与第二个代码一致的导入风格
from .convert import convert_wav_to_pcm, generate_feedback_history, generate_feedback_test
from .xf_recognizer import recognize_pcm
from .speech_evaluation import (
    clean_recognition_result,
    evaluate_text,
    generate_feedback,
    evaluate_intro_text,
    generate_feedback_intro
)
from .xf_spark_api import call_spark_x1

# 确保中文显示正常
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]
plt.rcParams['axes.unicode_minus'] = False


# 初始化会话状态（与第二个代码保持一致）
def _init_audio_session_state():
    if 'interview_questions' not in st.session_state:
        st.session_state.interview_questions = []
    if 'generated_intro' not in st.session_state:
        st.session_state.generated_intro = None
    if 'history_questions' not in st.session_state:
        st.session_state.history_questions = []
    if 'practice_history' not in st.session_state:
        st.session_state.practice_history = {}
    if 'interview_qa' not in st.session_state:
        st.session_state.interview_qa = {}


# 核心功能入口
def show_audio_app():
    _init_audio_session_state()

    # 讯飞API配置（与第二个代码一致）
    APPID = "79be4290"
    APIKey = "867eb662349b45edf64a4d48bc638a62"
    APISecret = "ZjE5NTZiYWFmNWNlMjg0OTUzMWVlM2Uz"
    DOMAIN = "x1"
    SPARK_URL = "wss://spark-api.xf-yun.com/v1/x1"

    # 文件夹配置（与第二个代码一致）
    UPLOAD_FOLDER = 'uploads'
    PCM_FOLDER = 'pcm_files'
    INTERVIEW_FOLDER = 'interviews'
    INTRO_FOLDER = 'self_intro_recordings'
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(PCM_FOLDER, exist_ok=True)
    os.makedirs(INTERVIEW_FOLDER, exist_ok=True)
    os.makedirs(INTRO_FOLDER, exist_ok=True)

    # 加载自定义CSS
    load_custom_css()

    # 页面标题
    st.markdown("""
        <h1 class='sub-header'>音频面试评测智能体</h1>
        """, unsafe_allow_html=True)

    # 侧边栏导航（核心修改：改回侧边栏）
    with st.sidebar:
        st.markdown("""
            <h5 class='sub-header'>音频功能导航</h3>
            """, unsafe_allow_html=True)
        # 侧边栏选项
        page = st.radio(
            "",  # 无标签，保持简洁
            ["上传音频分析", "模拟面试场景", "生成自我介绍"]
        )

        # 帮助按钮逻辑
        def toggle_help():
            st.session_state.show_help = not st.session_state.get('show_help', False)

        st.button("使用帮助", on_click=toggle_help)

        # 显示帮助信息
        if st.session_state.get('show_help', False):
            st.markdown("""
                ### 音频功能帮助
                - **上传音频分析**：支持WAV/PCM格式，大小不超过10MB
                - **模拟面试场景**：生成问题后需录制音频回答
                - **生成自我介绍**：填写信息后可生成文本并录音练习
                """)

    # 页面渲染分发（根据侧边栏选择显示对应功能）
    if page == "上传音频分析":
        _show_upload_audio_analysis(
            UPLOAD_FOLDER, PCM_FOLDER, INTERVIEW_FOLDER,
            APPID, APIKey, APISecret
        )
    elif page == "模拟面试场景":
        _show_simulation_interview(
            INTERVIEW_FOLDER, PCM_FOLDER,
            APPID, APIKey, APISecret, DOMAIN, SPARK_URL
        )
    elif page == "生成自我介绍":
        _show_self_intro_generation(
            INTRO_FOLDER, PCM_FOLDER,
            APPID, APIKey, APISecret, DOMAIN, SPARK_URL
        )

    # 页脚
    st.markdown("""
        ---
        <div style="text-align: center; color: #666; font-size: 0.9rem;">
            音频面试评测智能体 | 基于讯飞API提供服务
        </div>
        """, unsafe_allow_html=True)

# 内部函数1：上传音频分析（修正参数和调用）
def _show_upload_audio_analysis(upload_folder, pcm_folder, interview_folder, appid, apikey, apisecret):
    # 功能1：上传新音频分析
    st.subheader("上传面试音频")
    upload_domain = st.text_input("技术领域（如：人工智能）", "通用", key="upload_domain")
    upload_position = st.text_input("岗位类型（如：技术岗）", "通用", key="upload_position")

    uploaded_file = st.file_uploader(
        "选择音频文件（WAV/PCM）",
        type=['wav', 'pcm'],
        key="new_interview_audio_upload"
    )

    if uploaded_file is not None:
        filename = uploaded_file.name
        filepath = os.path.join(upload_folder, filename)
        with open(filepath, 'wb') as f:
            f.write(uploaded_file.getbuffer())

        ext = os.path.splitext(filename)[-1].lower()
        pcm_path = None
        if ext == '.pcm':
            pcm_path = filepath
        elif ext == '.wav':
            pcm_name = os.path.splitext(filename)[0] + '.pcm'
            pcm_path = os.path.join(pcm_folder, pcm_name)
            try:
                with st.spinner('正在转换音频格式...'):
                    convert_wav_to_pcm(filepath, pcm_path)
            except Exception as e:
                st.error(f"格式转换失败：{e}")

        if pcm_path and os.path.exists(pcm_path):
            with st.spinner('正在分析回答...'):
                with st.spinner('正在进行语音识别...'):
                    # 修正recognize_pcm调用
                    raw_result = recognize_pcm(appid, apikey, apisecret, pcm_path)
                result = clean_recognition_result(raw_result)

                scores = evaluate_text(
                    text=result,
                    domain=upload_domain,
                    position=upload_position
                )

                # 修正generate_feedback调用：移除多余参数
                feedback = generate_feedback(
                    scores,
                    result,
                    domain=upload_domain,
                    position=upload_position
                )

            st.subheader('📝 识别文本：')
            st.write(result)
            st.subheader('📊 能力维度评分：')
            for item, score in scores.items():
                st.write(f"{item}：{score} 分")

            st.subheader('💡 智能反馈建议：')
            if feedback:
                lines = [line.strip() for line in feedback.split('\n') if line.strip()]
                comment_lines = []
                suggestion_lines = []
                current_part = None

                for line in lines:
                    if " 内容点评" in line:
                        current_part = "comment"
                    elif " 改进建议" in line:
                        current_part = "suggestion"
                    elif current_part == "comment":
                        comment_lines.append(line)
                    elif current_part == "suggestion":
                        suggestion_lines.append(line)

                # 前端在显示前添加：
                st.markdown(" 内容点评")
                if comment_lines:
                    for line in comment_lines:
                        st.markdown(f"- {line}")
                else:
                    st.markdown("- 暂未生成内容点评")  # 空内容时显示默认文本

                st.markdown(" 改进建议")
                if suggestion_lines:
                    for line in suggestion_lines:
                        st.markdown(f"- {line}")
                else:
                    st.markdown("- 暂未生成改进建议")

    st.subheader("练习历史面试问题")

    with st.expander("📜 历史面试问题", expanded=False):
        history_domain = st.text_input("技术领域", "通用", key="history_domain")
        history_position = st.text_input("岗位类型", "通用", key="history_position")

        if os.path.exists(interview_folder):
            history_files = [f for f in os.listdir(interview_folder) if f.endswith(".json")]
            if history_files:
                col_sel, col_load = st.columns(2)
                with col_sel:
                    selected_file = st.selectbox("选择历史记录", history_files, key="history_file_sel")
                with col_load:
                    load_questions = st.button("提取历史问题", key="load_history_questions")

                if load_questions:
                    file_path = os.path.join(interview_folder, selected_file)
                    with open(file_path, "r", encoding="utf-8") as f:
                        record = json.load(f)
                    history_questions = list({ans["question"] for ans in record["answers"]})
                    st.session_state.history_questions = history_questions
                    st.success(f"已提取 {len(history_questions)} 个历史问题")

                if "history_questions" in st.session_state and st.session_state.history_questions:
                    selected_question = st.selectbox(
                        "选择历史问题",
                        st.session_state.history_questions,
                        key="selected_history_question"
                    )

                    if selected_question:
                        st.markdown("#### 选中的历史问题：")
                        st.info(selected_question)

                    new_audio = st.file_uploader(
                        "上传针对上述问题的新回答（WAV/PCM）",
                        type=['wav', 'pcm'],
                        key=f"history_audio_{hash(selected_question) % 10000}"
                    )

                    if new_audio is not None:
                        filename = f"history_{hash(selected_question) % 10000}_{new_audio.name}"
                        filepath = os.path.join(upload_folder, filename)
                        with open(filepath, 'wb') as f:
                            f.write(new_audio.getbuffer())

                        ext = os.path.splitext(filename)[-1].lower()
                        pcm_path = None
                        if ext == '.pcm':
                            pcm_path = filepath
                        elif ext == '.wav':
                            pcm_name = os.path.splitext(filename)[0] + '.pcm'
                            pcm_path = os.path.join(pcm_folder, pcm_name)
                            try:
                                with st.spinner('正在转换音频格式...'):
                                    convert_wav_to_pcm(filepath, pcm_path)
                            except Exception as e:
                                st.error(f"格式转换失败：{e}")

                        if pcm_path and os.path.exists(pcm_path):
                            with st.spinner('正在分析回答...'):
                                # 修正recognize_pcm调用：使用位置参数，与第二个代码一致
                                raw_answer = recognize_pcm(appid, apikey, apisecret, pcm_path)
                                clean_answer = clean_recognition_result(raw_answer)

                                scores = evaluate_text(
                                    text=clean_answer,
                                    domain=history_domain,
                                    position=history_position
                                )

                                feedback = generate_feedback_history(scores, clean_answer)

                            st.subheader("📊 新回答分析结果")
                            st.write("识别文本：", clean_answer)
                            st.subheader('📊 智能评分：')
                            for item, score in scores.items():
                                st.write(f"{item}：{score} 分")
                            st.subheader('💡 详细反馈：')
                            st.markdown("**🌟 能力点评**")
                            for item in feedback["🌟 能力点评"]:
                                st.markdown(f"- {item}")
                            st.markdown("**📈 改进建议**")
                            for item in feedback["📈 改进建议"]:
                                st.markdown(f"- {item}")
                else:
                    st.info("暂无历史记录，分析并保存面试后可在此查看")
            else:
                st.info("暂无历史记录文件夹，保存分析结果后自动创建")


# 内部函数2：模拟面试场景（修正参数和调用）
# 前端展示适配（Streamlit）
import streamlit as st
import os
from typing import Dict, List


def render_feedback(formatted_feedback: List):
    """在前端渲染结构化反馈，章节标题更小，避免重复展示"""
    if not formatted_feedback:
        st.warning("无反馈内容")
        return

    processed_sections = set()
    sections = ["回答正确性", "能力点评", "改进建议", "推荐资源"]

    for item in formatted_feedback:
        item_clean = item.strip()
        if not item_clean:
            continue

        # 章节标题用 ### 层级（比 subheader 小），只展示一次
        if item_clean in sections:
            if item_clean not in processed_sections:
                processed_sections.add(item_clean)
                # 使用 markdown 三级标题，字体更小
                st.markdown(f"### {item_clean}")
        elif processed_sections:  # 只展示已识别章节的内容
            try:
                if '.' in item_clean:
                    num, content = item_clean.split('.', 1)
                    st.markdown(f"- {num}. {content.strip()}")  # 序号内容用列表样式
                else:
                    st.markdown(f"- {item_clean}")  # 非序号内容也用列表
            except:
                st.markdown(f"- {item_clean}")  # 异常情况保底展示

def _show_simulation_interview(interview_folder, pcm_folder, appid, apikey, apisecret, domain, spark_url):
    st.subheader("选择面试场景并上传回答录音")
    domain = st.selectbox("选择技术领域", ["人工智能", "大数据", "物联网"])
    position = st.selectbox("选择岗位类型", ["技术岗", "产品岗", "运维测试岗"])

    # 生成面试问题
    if st.button("生成面试问题"):
        with st.spinner("正在生成面试问题..."):
            prompt = f"""
                请生成2道{domain}领域{position}的面试问题，要求：
                1. 覆盖专业核心知识；
                2. 包含1道实际项目问题；
                3. 问题清晰简洁，适合模拟面试。
                输出格式：每行1个问题，无需序号。
                """
            try:
                questions_str = call_spark_x1(prompt)
                questions = [q.strip() for q in questions_str.split("\n") if q.strip()]

                if len(questions) < 1:
                    st.warning("未生成有效问题，请重试")
                else:
                    st.session_state.interview_qa = {}  # 清空旧数据
                    # 限制最多3个问题，避免索引问题
                    for i, q in enumerate(questions[:3], 1):
                        st.session_state.interview_qa[f"question_{i}"] = {
                            "question": q,
                            "answer_text": "",
                            "scores": {},
                            "feedback": [],  # 初始化空反馈列表
                            "answered": False,
                            "audio_file": ""
                        }
                    st.success(f"成功生成 {len(questions[:3])} 个面试问题！")
            except Exception as e:
                st.error(f"生成问题失败：{str(e)}")
                if st.checkbox("显示详细错误信息"):
                    st.code(str(e), language="text")

    # 显示面试问题并处理回答
    if hasattr(st.session_state, 'interview_qa') and st.session_state.interview_qa:
        st.markdown(f"### {domain} - {position} 面试问题")

        question_keys = list(st.session_state.interview_qa.keys())
        for idx, q_key in enumerate(question_keys):
            qa = st.session_state.interview_qa[q_key]
            q_num = idx + 1
            st.markdown(f"**问题 {q_num}:** {qa['question']}")

            # 上传回答录音
            st.subheader(f"上传问题 {q_num} 的回答录音")
            answer_file = st.file_uploader(
                f"选择问题 {q_num} 的回答录音（WAV/PCM）",
                type=['wav', 'pcm'],
                key=f"answer_{q_key}"
            )

            if answer_file is not None:
                # 保存音频文件
                ans_filename = f"q{q_num}_{answer_file.name}"
                ans_filepath = os.path.join(interview_folder, ans_filename)
                with open(ans_filepath, 'wb') as f:
                    f.write(answer_file.getbuffer())

                # 格式转换
                ans_ext = os.path.splitext(ans_filename)[-1].lower()
                if ans_ext == '.pcm':
                    ans_pcm_path = ans_filepath
                elif ans_ext == '.wav':
                    ans_pcm_name = os.path.splitext(ans_filename)[0] + '.pcm'
                    ans_pcm_path = os.path.join(pcm_folder, ans_pcm_name)
                    try:
                        with st.spinner('正在转换回答录音为PCM...'):
                            convert_wav_to_pcm(ans_filepath, ans_pcm_path)  # 假设已实现此函数
                    except Exception as e:
                        st.error(f"回答录音转换失败：{e}")
                        continue
                else:
                    st.error("只支持上传 WAV 或 PCM 格式的录音")
                    continue

                # 语音识别与分析
                if os.path.exists(ans_pcm_path):
                    with st.spinner('正在分析回答内容...'):
                        try:
                            # 语音识别
                            raw_answer = recognize_pcm(appid, apikey, apisecret, ans_pcm_path)  # 假设已实现
                            clean_answer = clean_recognition_result(raw_answer)  # 假设已实现

                            # 能力评分
                            scores = evaluate_text(
                                text=clean_answer,
                                domain=domain,
                                position=position
                            )  # 假设已实现
                            if not scores:
                                scores = {"专业知识水平": 0}

                            # 生成反馈
                            formatted_feedback = generate_feedback_test(
                                scores=scores,
                                text=clean_answer,
                                domain=domain,
                                position=position,
                                question=qa['question']
                            )

                            # 保存结果
                            st.session_state.interview_qa[q_key].update({
                                "answer_text": clean_answer,
                                "scores": scores,
                                "feedback": formatted_feedback,
                                "answered": True,
                                "audio_file": ans_filename
                            })

                            # 前端展示（只展示一次）
                            st.success(f"问题 {q_num} 分析完成！")

                            # 显示面试分析结果标题（只显示一次）
                            st.markdown("### 📊 面试回答分析结果")
                            st.markdown(f"**问题 {q_num}:** {qa['question']}")
                            st.markdown("**📝 回答文本：**")
                            st.write(clean_answer)

                            # 显示评分
                            st.markdown("**智能评分（基于AI评估）：**")
                            score_cols = st.columns(len(scores))
                            for col, (name, score) in zip(score_cols, scores.items()):
                                col.metric(name, f"{score} 分")

                            # 显示详细反馈（通过render_feedback确保只展示一次）
                            st.markdown("**💡 详细反馈：**")
                            render_feedback(formatted_feedback)

                        except Exception as e:
                            st.error(f"处理失败：{str(e)}")

    # 保存记录（移到循环外部，确保只渲染一次）
    if st.session_state.interview_qa:
        st.subheader("💾 保存面试分析记录")
        default_name = f"interview_{domain}_{position}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        save_name = st.text_input("请输入保存名称（无需后缀）", value=default_name)
        if st.button("确认保存", key="save_all_interviews"):
            if not save_name.strip():
                st.error("保存名称不能为空！")
            else:
                os.makedirs(interview_folder, exist_ok=True)
                record_path = os.path.join(interview_folder, f"{save_name}.json")
                # 构建记录内容
                record = {
                    "domain": domain,
                    "position": position,
                    "questions": [qa['question'] for qa in st.session_state.interview_qa.values()],
                    "answers": [
                        {
                            "question": qa['question'],
                            "answer_text": qa['answer_text'],
                            "scores": qa['scores'],
                            "feedback": qa['feedback'],
                            "audio_file": qa.get('audio_file', '')
                        }
                        for qa in st.session_state.interview_qa.values() if qa['answered']
                    ],
                    "timestamp": datetime.now().strftime("%Y%m%d_%H%M%S")
                }
                with open(record_path, "w", encoding="utf-8") as f:
                    json.dump(record, f, ensure_ascii=False, indent=2)
                st.success(f"记录已保存至：{record_path}")


# 内部函数3：生成自我介绍（修正参数和调用）
def _show_self_intro_generation(intro_folder, pcm_folder, appid, apikey, apisecret, domain, spark_url):
    st.subheader("生成并上传自我介绍录音")
    st.info("步骤：1. 生成针对性自我介绍 → 2. 上传练习录音 → 3. 获取AI评分与改进建议")

    # 初始化变量
    eval_ok = False
    clean_intro = ""
    scores = {}
    feedback = ""

    # 1. 生成自我介绍
    with st.expander("填写个人信息（生成自我介绍用）", expanded=True):
        name = st.text_input("您的姓名（必填）", "")
        age = st.text_input("年龄（选填）", "")
        education = st.text_input("学历（必填）", "")
        major = st.text_input("专业（必填）", "例：计算机科学与技术")
        school = st.text_input("毕业院校（必填）", "例：XX大学")
        experience = st.text_area(
            "工作/实习经历（按STAR法则填写，未填自动生成相关）",
            "例：在XX公司实习期间，负责XX项目的XX模块开发..."
        )
        project = st.text_area(
            "核心项目经历（选填）",
            "例：主导XX课程设计，设计并实现XX系统..."
        )
        skills = st.multiselect(
            "专业技能（最多选5项）",
            [
                "机器学习", "深度学习", "计算机视觉", "自然语言处理", "TensorFlow", "PyTorch",
                "Hadoop", "Spark", "SQL", "数据挖掘", "数据分析", "数据仓库",
                "嵌入式系统", "传感器技术", "无线通信", "边缘计算", "云平台开发"
            ]
        )
        career_goal = st.text_input("职业目标（选填）", "例：希望从事算法研发工作")
        apply_reason = st.text_input("应聘动机（选填）", "例：对大数据技术有浓厚兴趣")
        domain = st.selectbox("目标技术领域", ["人工智能", "大数据", "物联网"])
        position = st.selectbox("目标岗位类型", ["技术岗", "产品岗", "运维测试岗"])
        # 生成自我介绍按钮
        if st.button("📝 生成自我介绍"):
            if not name:
                st.warning("请输入姓名（自我介绍的核心要素）")
            else:
                with st.spinner("正在生成个性化自我介绍..."):
                    prompt = f"""
                                请根据以下信息生成{domain}{position}{age}面试用的自我介绍（200-300字），结构需包含：
                                1. 基础信息：姓名{name}、学历（{education} {major}）、毕业院校（{school}）；
                                2. 核心经历：简述工作/实习经历（{experience}）或项目经历（{project}）；
                                3. 技能匹配：结合{domain}{position}需求，强调核心技能（{', '.join(skills[:5])}）；
                                4. 求职动机：选择该领域的原因（{apply_reason}）及短期目标（{career_goal}）。
                                要求：口语化、突出匹配度、时长60-90秒。
                                """
                    try:
                        intro = call_spark_x1(prompt)
                        st.session_state.generated_intro = intro
                        st.success("自我介绍生成成功！")
                    except Exception as e:
                        st.error(f"生成失败：{str(e)}")
                        if st.checkbox("显示详细错误"):
                            st.code(str(e))

        # 显示生成的自我介绍
        if "generated_intro" in st.session_state and st.session_state.generated_intro:
            if st.session_state.generated_intro.strip():
                st.markdown("### 📋 生成的自我介绍")
                st.write(st.session_state.generated_intro)  # 直接显示文本，不使用代码块
                st.info("✅ 包含要素：姓名学历 | 核心经历 | 技能匹配 | 求职动机")
            else:
                st.warning("生成的自我介绍为空，请重试")
        else:
            st.info("请先生成自我介绍")
        # 生成DOC文档并提供下载（修改为左对齐）
        def create_intro_docx(content):
            """创建符合格式要求的DOC文档：首行缩进两格、宋体四号字体、左对齐"""
            doc = Document()

            # 设置全局字体（宋体）
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(14)  # 四号字体对应14磅

            # 设置中文字体支持（避免宋体显示异常）
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 添加内容段落
            para = doc.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐

            # 设置首行缩进（两格，对应32磅，1格=16磅）
            para_format = para.paragraph_format
            para_format.first_line_indent = Pt(32)

            # 添加文本内容
            para.add_run(content)

            # 保存到内存流
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        # 生成DOC文件并提供下载按钮
        doc_buffer = create_intro_docx(st.session_state.generated_intro)
        st.download_button(
            label="📥 下载DOC格式自我介绍",
            data=doc_buffer,
            file_name=f"自我介绍_{domain}_{position}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_intro_docx"
        )

        # 原有保存文本功能（修改保存路径至results文件夹）
        if st.button("💾 保存TXT格式文本", key="save_intro_txt"):
            os.makedirs("results", exist_ok=True)  # 确保results文件夹存在
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"self_intro_{domain}_{position}_{timestamp}.txt"
            filepath = os.path.join("results", filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(st.session_state.generated_intro)
            st.success(f"已保存至：{filepath}")

        # 2. 上传自我介绍录音分析
        st.subheader("上传自我介绍练习录音")
        intro_file = st.file_uploader("选择录音文件（WAV/PCM）", type=['wav', 'pcm'], key="intro_upload")

        if intro_file is not None:
            # 保存上传的录音
            intro_filename = intro_file.name
            intro_filepath = os.path.join(intro_folder, intro_filename)
            with open(intro_filepath, 'wb') as f:
                f.write(intro_file.getbuffer())

            # 格式转换（WAV→PCM）
            intro_ext = os.path.splitext(intro_filename)[-1].lower()
            intro_pcm_path = None  # 初始化PCM路径变量
            if intro_ext == '.pcm':
                intro_pcm_path = intro_filepath
            elif intro_ext == '.wav':
                intro_pcm_name = os.path.splitext(intro_filename)[0] + '.pcm'
                intro_pcm_path = os.path.join(pcm_folder, intro_pcm_name)
                try:
                    with st.spinner('正在转换录音为PCM...'):
                        convert_wav_to_pcm(intro_filepath, intro_pcm_path)
                except Exception as e:
                    st.error(f"录音转换失败：{e}")
            else:
                st.error("只支持上传 WAV 或 PCM 格式的录音")

            # 语音识别+分析（针对自我介绍）
            if intro_pcm_path and os.path.exists(intro_pcm_path):
                with st.spinner('正在分析自我介绍内容...'):
                    try:
                        raw_intro = recognize_pcm(appid, apikey, apisecret, intro_pcm_path)
                        clean_intro = clean_recognition_result(raw_intro)
                        recog_ok = True
                    except Exception as e:
                        st.error(f"自我介绍识别失败：{e}")
                        recog_ok = False
                        clean_intro = "识别失败，无内容"

                if recog_ok:
                    try:
                        scores = evaluate_intro_text(
                            text=clean_intro,
                            domain=domain,
                            position=position
                        )
                        feedback = generate_feedback_intro(
                            scores=scores,
                            text=clean_intro,
                            target_position=f"{domain}{position}",
                            ideal_version=st.session_state.generated_intro
                        )
                        eval_ok = True
                    except Exception as e:
                        st.error(f"评分失败：{e}")
                        eval_ok = False

                if eval_ok:
                    # 1. 自我介绍专项评分展示
                    st.markdown("### 📊 自我介绍专项评分（满分100分）")
                    for dim, score in scores.items():
                        # 结合状态标签，清晰区分高低分
                        if score < 60:
                            st.markdown(f"**{dim}：{score} 分** ❌（需重点优化）")
                        elif score >= 80:
                            st.markdown(f"**{dim}：{score} 分** ✅（优势保持）")
                        else:
                            st.markdown(f"**{dim}：{score} 分** ⚠️（待提升）")

                    # 2. 内容对比展示
                    st.subheader('📝 内容对比（与理想版本）')
                    col_actual, col_ideal = st.columns(2)
                    with col_actual:
                        st.markdown("**您的录音内容：**")
                        st.write(clean_intro or "未识别到有效内容")  # 容错：无内容时显示提示
                    with col_ideal:
                        st.markdown("**理想版本参考：**")
                        st.write(st.session_state.generated_intro or "未生成理想版本")  # 容错处理

                    # 3. 改进建议展示（核心逻辑）
                    st.subheader('💡 改进建议')

                    # # 调试信息：展示后端传递的原始反馈（方便定位问题）
                    # with st.expander("调试信息（原始反馈）", expanded=False):
                    #     st.code(feedback, language="text")  # 直接显示后端传递的原始字符串

                    # 改进建议展示（精简版）
                    if not feedback:
                        st.warning("未生成有效改进建议，请重试")
                    else:
                        # 基础分割：按换行拆分并过滤空行
                        feedback_lines = [line.strip() for line in feedback.split('\n') if line.strip()]

                        # 尝试结构化解析
                        feedback_items = []
                        current_item = ""
                        for line in feedback_lines:
                            # 匹配以数字序号开头的行（兼容中英文句号和空格）
                            if re.match(r'^\s*\d+\s*[.．]\s*', line):
                                if current_item:  # 保存上一条建议
                                    feedback_items.append(current_item)
                                current_item = line  # 开始新建议
                            else:
                                # 合并同序号的多行内容
                                if current_item:
                                    current_item += f" {line}"
                                else:
                                    current_item = line  # 无序号的行作为新建议起始

                        # 保存最后一条建议
                        if current_item:
                            feedback_items.append(current_item)

                        # 过滤空项
                        feedback_items = [item for item in feedback_items if item.strip()]

                        # 层级展示：核心建议 + 完整建议
                        if feedback_items:
                            # 核心建议（前2条）
                            st.markdown("**核心建议**（优先改进）：")
                            for item in feedback_items[:2]:
                                # 清理序号格式后展示
                                clean_item = re.sub(r'^\s*\d+\s*[.．]\s*', '', item)
                                st.markdown(f"- {clean_item}")

                            # 完整建议（带序号修正）
                            with st.expander("查看完整建议", expanded=False):
                                for i, item in enumerate(feedback_items, 1):
                                    # 统一序号格式（替换中文句号为英文句号）
                                    clean_item = item.replace('．', '.')
                                    # 修正序号错误，确保按1.2.3顺序排列
                                    if re.match(r'^\d+\.', clean_item):
                                        content = re.sub(r'^\d+\.\s*', '', clean_item)
                                        st.markdown(f"{i}. {content}")
                                    else:
                                        st.markdown(f"{i}. {clean_item}")
                        else:
                            # 无有效分割项时直接展示原始反馈
                            st.markdown("**建议内容：**")
                            st.markdown(feedback)
                            with st.expander("查看完整建议"):
                                st.markdown(feedback)

if __name__ == "__main__":
    show_audio_app()
