import streamlit as st
import os
import io
import tempfile
import subprocess
import base64
import pythoncom  # 导入 pythoncom 用于 COM 库初始化


class ResumeEditor:
    @staticmethod
    def initialize_resume():
        """初始化空白简历数据结构"""
        return {
            "name": "",
            "contact": {"email": "", "phone": ""},
            "education": [{"degree": "", "school": "", "time": ""}],
            "experience": [{"company": "", "position": "", "time": "", "description": ""}],
            "projects": [{"name": "", "role": "", "time": "", "description": ""}],
            "skills": []  # 添加技能字段
        }


def create_docx(resume):
    from docx import Document
    doc = Document()
    doc.add_heading("个人简历")
    doc.add_heading('基本信息', level=1)
    doc.add_paragraph(f"姓名: {resume['name']}")
    doc.add_paragraph(f"邮箱: {resume['contact']['email']}")
    doc.add_paragraph(f"电话: {resume['contact']['phone']}")

    doc.add_heading('教育经历', level=1)
    for edu in resume["education"]:
        doc.add_paragraph(f"{edu['degree']} - {edu['school']} ({edu['time']})")

    doc.add_heading('工作经验', level=1)
    for exp in resume["experience"]:
        doc.add_paragraph(f"{exp['position']} - {exp['company']} ({exp['time']})")
        doc.add_paragraph(exp["description"])

    doc.add_heading('项目经历', level=1)
    for proj in resume["projects"]:
        doc.add_paragraph(f"{proj['role']} - {proj['name']} ({proj['time']})")
        doc.add_paragraph(proj["description"])

    doc.add_heading('技能描述', level=1)
    for skill in resume["skills"]:
        doc.add_paragraph(skill)

    return doc


def convert_to_pdf(docx_path, pdf_path):
    try:
        if os.name == 'nt':  # Windows
            pythoncom.CoInitialize()  # 初始化 COM 库
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False  # 不显示 Word 窗口
            try:
                if not os.path.exists(docx_path):
                    st.error(f"临时 DOCX 文件 {docx_path} 不存在，请检查。")
                    return
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(pdf_path, FileFormat=17)
                doc.Close(SaveChanges=0)  # 关闭文档，不保存更改
            except Exception as e:
                if isinstance(e.args, tuple) and e.args[0] == -2147023170:
                    return ""
                st.error(f"打开或保存文件时出错: {e}")
            finally:
                word.Quit()  # 退出 Word 应用程序
                pythoncom.CoUninitialize()  # 释放 COM 库资源
        else:  # Linux
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path),
                            docx_path])
    except Exception as e:
        st.error(f"转换为 PDF 时出错: {e}")


def resume_management():
    st.title("在线简历管理")
    user = st.session_state.current_user

    # 初始化简历数据（首次使用或无数据时）
    if 'edited_resume' not in st.session_state:
        # 尝试从用户数据加载已有简历，无则初始化空白简历
        user_data = st.session_state.user_manager.users.get(user, {})
        existing_resume = user_data.get("parsed_resume")

        if existing_resume:
            # 确保简历数据包含所有必要字段，特别是新添加的 skills 字段
            resume_template = ResumeEditor.initialize_resume()
            st.session_state.edited_resume = {**resume_template, **existing_resume}

            # 确保嵌套结构正确
            for key in ['education', 'experience', 'projects']:
                if key not in st.session_state.edited_resume or not isinstance(
                        st.session_state.edited_resume[key],
                        list):
                    st.session_state.edited_resume[key] = resume_template[key]
        else:
            st.session_state.edited_resume = ResumeEditor.initialize_resume()

    # 基本信息编辑
    st.subheader("基本信息")
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.edited_resume["name"] = st.text_input("姓名", st.session_state.edited_resume["name"])
    with col2:
        st.session_state.edited_resume["contact"]["email"] = st.text_input(
            "邮箱", st.session_state.edited_resume["contact"]["email"]
        )

    col3, col4 = st.columns(2)
    with col3:
        st.session_state.edited_resume["contact"]["phone"] = st.text_input(
            "电话", st.session_state.edited_resume["contact"]["phone"]
        )

    # 教育经历编辑
    st.subheader("教育经历")
    edu_list = st.session_state.edited_resume["education"]

    # 添加新教育经历
    if st.button("➕ 添加教育经历"):
        edu_list.append({"degree": "", "school": "", "time": ""})
        st.rerun()

    # 编辑现有教育经历
    for i in range(len(edu_list) - 1, -1, -1):
        edu = edu_list[i]
        with st.expander(f"教育经历 {i + 1}", expanded=True):
            col1, col2 = st.columns([3, 1])
            with col1:
                edu["school"] = st.text_input(f"学校 {i + 1}", edu["school"], key=f"edu_school_{i}")
                edu["degree"] = st.text_input(f"学位 {i + 1}", edu["degree"], key=f"edu_degree_{i}")
                edu["time"] = st.text_input(f"时间（如：2018-2022）", edu["time"], key=f"edu_time_{i}")
            with col2:
                if st.button(f"🗑️ 删除", key=f"del_edu_{i}"):
                    del edu_list[i]
                    st.success(f"已删除教育经历 {i + 1}")
                    st.rerun()

    # 工作经验编辑
    st.subheader("工作经验")
    exp_list = st.session_state.edited_resume["experience"]

    # 添加新工作经验
    if st.button("➕ 添加工作经验"):
        exp_list.append({"company": "", "position": "", "time": "", "description": ""})
        st.rerun()

    # 编辑现有工作经验
    for i in range(len(exp_list) - 1, -1, -1):
        exp = exp_list[i]
        with st.expander(f"工作经验 {i + 1}", expanded=True):
            col1, col2 = st.columns([3, 1])
            with col1:
                exp["company"] = st.text_input(f"公司名称 {i + 1}", exp["company"], key=f"exp_company_{i}")
                exp["position"] = st.text_input(f"职位 {i + 1}", exp["position"], key=f"exp_position_{i}")
                exp["time"] = st.text_input(f"工作时间", exp["time"], key=f"exp_time_{i}")
                exp["description"] = st.text_area(
                    f"工作描述（职责与成就）", exp["description"], key=f"exp_desc_{i}", height=100
                )
            with col2:
                if st.button(f"🗑️ 删除", key=f"del_exp_{i}"):
                    del exp_list[i]
                    st.success(f"已删除工作经验 {i + 1}")
                    st.rerun()

    # 项目经历编辑
    st.subheader("项目经历")
    project_list = st.session_state.edited_resume["projects"]

    # 添加新项目经历
    if st.button("➕ 添加项目经历"):
        project_list.append({"name": "", "role": "", "time": "", "description": ""})
        st.rerun()

    # 编辑现有项目经历
    for i in range(len(project_list) - 1, -1, -1):
        project = project_list[i]
        with st.expander(f"项目经历 {i + 1}", expanded=True):
            col1, col2 = st.columns([3, 1])
            with col1:
                project["name"] = st.text_input(f"项目名称 {i + 1}", project["name"], key=f"proj_name_{i}")
                project["role"] = st.text_input(f"担任角色", project["role"], key=f"proj_role_{i}")
                project["time"] = st.text_input(f"项目时间", project["time"], key=f"proj_time_{i}")
                project["description"] = st.text_area(
                    f"项目描述（职责与成果）", project["description"], key=f"proj_desc_{i}", height=100
                )
            with col2:
                if st.button(f"🗑️ 删除", key=f"del_proj_{i}"):
                    del project_list[i]
                    st.success(f"已删除项目经历 {i + 1}")
                    st.rerun()

    # 技能编辑
    st.subheader("技能描述")
    skills_list = st.session_state.edited_resume["skills"]

    # 添加新技能
    if st.button("➕ 添加技能"):
        skills_list.append("")
        st.rerun()

    # 编辑现有技能
    for i in range(len(skills_list) - 1, -1, -1):
        skill = skills_list[i]
        col1, col2 = st.columns([3, 1])
        with col1:
            skills_list[i] = st.text_input(f"技能 {i + 1}", skill, key=f"skill_{i}")
        with col2:
            if st.button(f"🗑️ 删除技能", key=f"del_skill_{i}"):
                del skills_list[i]
                st.success(f"已删除技能 {i + 1}")
                st.rerun()

    # 保存功能
    col_save, col_reset = st.columns(2)
    with col_save:
        if st.button("💾 保存简历", type="primary"):
            success = st.session_state.user_manager.update_parsed_resume(user, st.session_state.edited_resume)
            if success:
                st.success("简历已成功保存！")
            else:
                st.error("保存失败，请重试")

    with col_reset:
        if st.button("🔄 重置为空白简历"):
            if st.warning("确定要清空所有内容吗？此操作不可恢复", icon="⚠️"):
                st.session_state.edited_resume = ResumeEditor.initialize_resume()
                st.rerun()

    # 预览和下载功能
    col_preview, col_download = st.columns(2)
    with col_preview:
        if st.button("📄 预览简历"):
            resume = st.session_state.edited_resume
            doc = create_docx(resume)
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)

            # 保存为临时文件
            temp_dir = tempfile.gettempdir()
            temp_docx_path = os.path.join(temp_dir, "temp_resume.docx")
            with open(temp_docx_path, 'wb') as f:
                f.write(docx_buffer.read())

            temp_pdf_path = os.path.join(temp_dir, "temp_resume.pdf")
            convert_to_pdf(temp_docx_path, temp_pdf_path)

            if os.path.exists(temp_pdf_path):
                with open(temp_pdf_path, "rb") as f:
                    pdf_bytes = f.read()
                pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
                st.markdown(
                    f'<embed src="data:application/pdf;base64,{pdf_base64}" width="100%" height="auto" type="application/pdf">',
                    unsafe_allow_html=True)
            else:
                st.error("无法生成 PDF 文件进行预览。")

    with col_download:
        if st.button("📥 下载简历"):
            resume = st.session_state.edited_resume
            doc = create_docx(resume)
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            st.download_button(
                label="下载 DOCX 文件",
                data=docx_buffer,
                file_name="resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
